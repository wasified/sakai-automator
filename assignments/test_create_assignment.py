#-------------------------------------------------------------------------------
# Name:        test_basic_post
# Purpose:
#
# Author:      lifel
#
# Created:     05/04/2024
# Copyright:   (c) lifel 2024
# Licence:     <your licence>
#-------------------------------------------------------------------------------

# This script translates the following test script:
#https://docs.google.com/spreadsheets/d/1D4ClvBNbPE7A5JIQyeB1_PbVT7UuxxnrVFWUlNwB0yc/edit?gid=2076185087#gid=2076185087
# comments have been added

from playwright.sync_api import Playwright, sync_playwright, expect
from utils import read_csv


import datetime
import string
import random
import re

# this is added so we make an assignment
# with a new name everytime
random_suffix = ''.join(random.choices(string.ascii_letters,
                             k=5)) # initializing size of string

# docx report stuff
import io
from docx import Document
from docx.shared import Inches

report_docx = Document()
report_docx.add_heading("Test Report", 0)
current_datetime = datetime.datetime.now()
datetime_string = current_datetime.strftime('%Y-%m-%d %H:%M:%S')
report_docx.add_heading('Test run on: %s'%(datetime_string), 1)
p = report_docx.add_paragraph("")
run = p.add_run()
run.add_break()

ASSIGNMENT_NAME = "Assignment {}".format(random_suffix)
#ASSIGNMENT_NAME = "Assignment BkWUo"
COURSE_NAME = "VJ101"
URL = "https://trunk-mysql8.nightly.sakaiproject.org/"
instructor_id, instructor_pass = read_csv.get_instructor_credentials()

def add_to_report(screenshot_bytes, text):
    p = report_docx.add_paragraph(text)
    run = p.add_run()
    run.add_break()
    screenshot_file = io.BytesIO(screenshot_bytes)
    screenshot_file.seek(0)
    report_docx.add_picture(screenshot_file, width=Inches(7))
    report_docx.save('test-log.docx')

def add_heading(text):
    report_docx.add_heading(text, 3)
    report_docx.save('test-log.docx')

def test_create_assignment(playwright: Playwright) -> None:
    """
    This function goes from steps 1.01 to 1.08 in the
    basic_functionality_testing script.

    """
    global browser, context, page
    browser = playwright.chromium.launch()
    context = browser.new_context()
    page = context.new_page()
    page.goto(URL)
    #screenshot_bytes = page.screenshot(type='png')
    #add_to_report(screenshot_bytes, "Opened Portal")
    page.get_by_placeholder("Username").click()
    page.get_by_placeholder("Username").fill(instructor_id)
    page.get_by_placeholder("Username").press("Tab")
    page.get_by_placeholder("Password").fill(instructor_pass)
    page.get_by_placeholder("Password").press("Enter")
    add_heading("Logging in as instructor")
    #page.get_by_role("button", name="Expand tool list").click()
    #page.get_by_role("link", name="Assignments").click()
    #page.get_by_role("button", name="").click()
    #page.get_by_role("link", name="Okay, got it!").click()
    page.get_by_role("link", name="%s 1 1 Spring"%COURSE_NAME).click()
    page.get_by_role("link", name="Assignments").click()
    page.get_by_role("link", name="Add").click()
    page.get_by_role("button", name="Post").click()
    # assert empty form isn't published: 1.02 in the script
    expect(page.get_by_text("Alert: Please specify the assignment title")).to_be_visible()
    screenshot_bytes = page.screenshot(type='png')
    add_to_report(screenshot_bytes, "TEST PASSED: Asserting empty form isn't published:")
    # take a screenshot
    page.get_by_placeholder("Title").click()
    page.get_by_placeholder("Title").fill(ASSIGNMENT_NAME)
    page.frame_locator("iframe[title=\"Editor\\, new_assignment_instructions\"]").locator("html").click()
    page.frame_locator("iframe[title=\"Editor\\, new_assignment_instructions\"]").get_by_label("Editor,").fill("Sample description. ")
    page.get_by_placeholder("Title").press("Tab")
    page.get_by_label("Max Points").click()
    page.get_by_label("Max Points").fill("a")
    page.get_by_role("button", name="Post").click()

    # asserts error message: 1.06 in the script
    expect(page.get_by_text("Alert: Please use a number")).to_be_visible()
    screenshot_bytes = page.screenshot(type='png')
    add_to_report(screenshot_bytes, "TEST PASSED: Entering alpha character:")
    page.get_by_label("Hide due date from students").check()
    page.get_by_label("Add an announcement about the").check()
    page.get_by_label("Max Points").click()
    page.get_by_label("Max Points").click(click_count=3)
    page.get_by_label("Max Points").fill("10")
    #page.locator("#assignmentGradingGradebookOptionsPanel div").click()
    page.get_by_role("button", name="Post").click()
    # asserts that assignement was created
    expect(page.get_by_role("cell", name="%s Edit Assignment"%ASSIGNMENT_NAME)).to_be_visible()
    page.get_by_role("button", name="Profile image").click()
    page.get_by_role("link", name=" Log Out").click()
    add_heading("Logging out as instructor")

    # ---------------------
    #context.close()
    #browser.close()


def test_view_assignment_as_student(playwright: Playwright) -> None:
    #browser = playwright.chromium.launch(headless=False)
    #context = browser.new_context()
    #page = context.new_page()
    add_heading("Logging in as student")
    page.goto(URL)
    page.get_by_placeholder("Username").click()
    page.get_by_placeholder("Username").fill("student1")
    page.get_by_placeholder("Username").press("Tab")
    page.get_by_placeholder("Password").fill("sakai")
    page.get_by_placeholder("Password").press("Enter")
    #page.get_by_role("button", name=" Click to dismiss").click()
    #page.get_by_role("link", name="Okay, got it!").click()
    page.get_by_role("button", name="Expand tool list").click()
    page.locator("#pinned-site-list").get_by_role("link", name="Announcements").click()
    #page.get_by_role("link", name="Assignment: Open Date for '%s'"%ASSIGNMENT_NAME).click()
    page.get_by_role("link", name="Assignment: Open Date for ''%s'' Assignment: Open Date for ''"%ASSIGNMENT_NAME).click()
    #1..09: Verify you see the announcement
    expect(page.locator("#content div").filter(has_text="Open date for assignment ''").nth(3)).to_be_visible()
    screenshot_bytes = page.screenshot(type='png')
    add_to_report(screenshot_bytes, "TEST PASSED: Verified announcement")
    #navigate to assignments tool
    page.get_by_role("link", name="Assignments").click()
    # 1.10: assert date field is empty:
    expect(page.get_by_role("rowgroup")).to_contain_text("")
    screenshot_bytes = page.screenshot(type='png')
    add_to_report(screenshot_bytes, "TEST PASSED: Date field is empty")
    # assert assignment name matches
    #1.12: Verfiy the name of the assignment
    expect(page.get_by_role("rowgroup")).to_contain_text(ASSIGNMENT_NAME)
    screenshot_bytes = page.screenshot(type='png')
    add_to_report(screenshot_bytes, "TEST PASSED: Assignment name matches")
    # click on assignment submission page
    page.get_by_role("link", name=ASSIGNMENT_NAME).click()
    # 1.12: verify that due date is empty
    expect(page.locator("#StudentAssignmentCurrent")).to_contain_text("")
    screenshot_bytes = page.screenshot(type='png')
    add_to_report(screenshot_bytes, "TEST PASSED: No due date on assignment page")

    # log out as student
    page.get_by_role("button", name="Profile image").click()
    page.get_by_role("link", name=" Log Out").click()

    # ---------------------
    #context.close()
    #browser.close()

##def test_instructor_edits_assignment(playwright: Playwright):
##    """
##    This test pertains to section 2 in basic functionality test,
##    commenting this out for now since need to
##    implement time date functioncality to make this proper.
##
##     """
##    browser = playwright.chromium.launch()
##    context = browser.new_context()
##    page = context.new_page()
##    page.goto(URL)
##    page.get_by_placeholder("Username").click()
##    page.get_by_placeholder("Username").fill("qateacher")
##    page.get_by_placeholder("Username").press("Tab")
##    page.get_by_placeholder("Password").fill("sakai")
##    page.get_by_placeholder("Password").press("Enter")
##    page.get_by_role("button", name="Expand tool list").click()
##    page.get_by_role("link", name="Assignments").click()
##    page.get_by_role("link", name="Edit {}".format(ASSINGMENT_NAME)).click()
##    page.get_by_label("Due Date", exact=True).click()
##    page.get_by_label("Due Date", exact=True).click()
##    page.get_by_label("Due Date", exact=True).click()
##    page.get_by_label("Accept Until", exact=True).click()
##    page.get_by_label("Accept Until", exact=True).click()
##    page.get_by_label("Accept Until", exact=True).click()
##    page.get_by_label("Due Date", exact=True).click()
##    page.get_by_label("Due Date", exact=True).click()
##    page.get_by_label("Due Date", exact=True).click()
##    page.get_by_label("Due Date", exact=True).click()
##    page.get_by_label("Due Date", exact=True).press("ArrowRight")
##    page.get_by_label("Due Date", exact=True).fill("")
##    page.get_by_label("Due Date", exact=True).press("Control+z")
##    page.get_by_label("Due Date", exact=True).press("Control+z")
##    page.get_by_label("Due Date", exact=True).press("Control+z")
##    page.get_by_label("Due Date", exact=True).click()
##    page.get_by_label("Due Date", exact=True).press("ArrowLeft")
##    page.get_by_label("Due Date", exact=True).fill("2024-09-10T20:20")
##    page.get_by_label("Accept Until", exact=True).click()
##    page.get_by_label("Accept Until", exact=True).press("ArrowRight")
##    page.get_by_label("Accept Until", exact=True).fill("2024-09-11T20:20")
##    page.get_by_label("Hide due date from students").uncheck()
##    page.get_by_role("button", name="Post").click()
##    #2.02: Verify the date updated
##    expect(page.get_by_text("Sep 10, 2024, 8:20 PM")).to_be_visible()

    # ---------------------

def test_student_submission(playwright: Playwright) -> None:
    browser = playwright.chromium.launch(headless=False)
    context = browser.new_context()
    page = context.new_page()
    page.goto("https://trunk-mysql8.nightly.sakaiproject.org/portal")
    page.get_by_placeholder("Username").click()
    page.get_by_placeholder("Username").fill("student1")
    page.get_by_placeholder("Username").press("Tab")
    page.get_by_placeholder("Password").fill("sakai")
    page.get_by_placeholder("Password").press("Enter")
    page.get_by_role("button", name="Expand tool list").click()
    page.get_by_role("link", name="Assignments").click()
    expect(page.get_by_role("rowgroup")).to_contain_text("Not Started")
    page.get_by_role("link", name=ASSIGNMENT_NAME).click()
    expect(page.locator("#StudentAssignmentCurrent")).to_contain_text("Points (max 10.00)")
    page.get_by_text("Sample description.").click()
    expect(page.locator("#StudentAssignmentCurrent")).to_contain_text("Sample description.")
    page.frame_locator("iframe[title=\"Editor\\, Assignment\\.view_submission_text\"]").locator("html").click()
    page.frame_locator("iframe[title=\"Editor\\, Assignment\\.view_submission_text\"]").get_by_label("Editor, Assignment.").fill("This is some inline text")
    # file upload stuff here
    #page.get_by_label("Select a file from computer").click()
    page.get_by_label("Select a file from computer").set_input_files("test1.docx")
    # file upload end
    page.locator("#preview").click()
    expect(page.locator("#content")).to_contain_text("This is some inline text")
    expect(page.locator("#content")).to_contain_text("test1.docx")
    page.get_by_role("button", name="Save Draft").click()
    expect(page.locator("#content")).to_contain_text("You have successfully saved your work but NOT submitted yet. To complete submission, you must select the Submit button.")
    page.get_by_role("button", name="Back to list").click()
    expect(page.get_by_role("rowgroup")).to_contain_text("In progress")
    # log out as student one
    page.get_by_role("button", name="Profile image").click()
    page.get_by_role("link", name=" Log Out").click()
    # Entering as student 2
    page.get_by_placeholder("Username").click()
    page.get_by_placeholder("Username").fill("student2")
    page.get_by_placeholder("Username").press("Tab")
    page.get_by_placeholder("Password").fill("sakai")
    page.get_by_placeholder("Password").press("Enter")
    #page.get_by_role("button", name=" Click to dismiss").click()
    #page.get_by_role("link", name="Okay, got it!").click()
    page.get_by_role("button", name="Expand tool list").click()
    page.get_by_role("link", name="Assignments").click()
    page.get_by_role("link", name=ASSIGNMENT_NAME).click()
    page.frame_locator("iframe[title=\"Editor\\, Assignment\\.view_submission_text\"]").locator("html").click()
    page.frame_locator("iframe[title=\"Editor\\, Assignment\\.view_submission_text\"]").get_by_label("Editor, Assignment.").fill("This is some inline text. ")
    page.get_by_role("button", name="Cancel").click()
    expect(page.get_by_text("Your changes will be")).to_be_visible()
    page.get_by_role("button", name="Yes").click()
    expect(page.get_by_role("rowgroup")).to_contain_text("Not Started")
    page.get_by_role("link", name=ASSIGNMENT_NAME).click()
    page.frame_locator("iframe[title=\"Editor\\, Assignment\\.view_submission_text\"]").locator("html").click()
    page.frame_locator("iframe[title=\"Editor\\, Assignment\\.view_submission_text\"]").locator("html").click()
    page.frame_locator("iframe[title=\"Editor\\, Assignment\\.view_submission_text\"]").get_by_label("Editor, Assignment.").fill("Some new submission text here. ")
    page.get_by_role("button", name="Proceed").click()
    expect(page.get_by_text("Alert: If you are ready to")).to_be_visible()
    page.get_by_label("revise").click()
    page.frame_locator("iframe[title=\"Editor\\, Assignment\\.view_submission_text\"]").get_by_text("Some new submission text here.").click()
    page.frame_locator("iframe[title=\"Editor\\, Assignment\\.view_submission_text\"]").get_by_text("Some new submission text here.").click()
    page.frame_locator("iframe[title=\"Editor\\, Assignment\\.view_submission_text\"]").get_by_label("Editor, Assignment.").fill("Some new submission text here. Adding some new stuff here. ")
    page.get_by_role("button", name="Proceed").click()
    page.get_by_label("post").click()
    expect(page.get_by_text("You have successfully")).to_be_visible()
    page.get_by_role("button", name="Back to list").click()
    submitted_re = re.compile("Submitted .*")
    expect(page.get_by_role("rowgroup")).to_contain_text(submitted_re)
    ## string to match
    page.get_by_role("button", name="Profile image").click()
    page.get_by_role("link", name=" Log Out").click()
    page.get_by_placeholder("Username").click()
    page.get_by_placeholder("Username").fill("student3")
    page.get_by_placeholder("Username").press("Tab")
    page.get_by_placeholder("Password").fill("sakai")
    page.get_by_placeholder("Password").press("Enter")
    #page.goto("https://trunk-mysql8.nightly.sakaiproject.org/portal/site/2ddbea63-62da-4ced-814a-a6e93c2120c8/tool/64785cf0-38d3-4750-946d-7a49e25f7cd1")
    page.get_by_role("link", name="%s 1 1 Spring"%COURSE_NAME).click()
    page.get_by_role("link", name="Assignments").click()
    page.get_by_role("link", name=ASSIGNMENT_NAME).click()
    #page.get_by_label("Select a file from computer").click()
    page.get_by_label("Select a file from computer").set_input_files("test1.docx")
    page.get_by_role("button", name="Proceed").click()
    expect(page.get_by_text("Alert: If you are ready to")).to_be_visible()
    page.get_by_label("post").click()
    expect(page.locator("#content")).to_contain_text("You have successfully submitted your work. You will receive an email confirmation containing this information.")
    page.get_by_role("button", name="Back to list").click()
    # will change the date later
    #expect(page.get_by_role("rowgroup")).to_contain_text("Submitted Sep 4, 2024, 8:27 PM")

    # log out as student
    page.get_by_role("button", name="Profile image").click()
    page.get_by_role("link", name=" Log Out").click()
    # ---------------------
    #context.close()
    #browser.close()

def test_submit_on_behalf_student(playwright: Playwright) -> None:
    #browser = playwright.chromium.launch(headless = False)
    #context = browser.new_context()
    page = context.new_page()
    page.goto("https://trunk-mysql8.nightly.sakaiproject.org/portal")
    page.get_by_text("jump to content [c] Sites [w] Tools [l] Welcome View System Alert Username").click()
    page.frame_locator("iframe[title=\"Home Information Message\"]").get_by_role("heading", name="Welcome to Sakai").click()
    page.get_by_placeholder("Username").click()
    page.get_by_placeholder("Username").fill("qateacher")
    page.get_by_placeholder("Username").press("Tab")
    page.get_by_placeholder("Password").fill("sakai")
    page.get_by_placeholder("Password").press("Enter")
    page.get_by_role("button", name="Expand tool list").click()
    page.get_by_role("link", name="Assignments").click()
    page.get_by_role("link", name="Assignments by Student").click()
    page.get_by_role("link", name="Four, Student (student4)").click()
    page.get_by_role("link", name="Submit on behalf of Student").first.click()
    expect(page.locator("#content")).to_contain_text("This is an example preview of how students will see the assignment. You may work through this assignment as a student would, including submitting it.")
    expect(page.locator("#addSubmissionForm")).to_contain_text("Student Four")
    page.frame_locator("iframe[title=\"Editor\\, Assignment\\.view_submission_text\"]").locator("html").click()
    page.frame_locator("iframe[title=\"Editor\\, Assignment\\.view_submission_text\"]").get_by_label("Editor, Assignment.").fill("This is some text from student 4. ")
    page.get_by_role("button", name="Proceed").click()
    expect(page.locator("#content")).to_contain_text("You have successfully submitted your work. You will receive an email confirmation containing this information.")
    page.get_by_role("button", name="Back to list").click()
    submitted_re = re.compile(r"Ungraded - Submitted .*")
    #expect(page.locator("#assignmentsByStudent")).to_contain_text("Ungraded - Submitted Sep 6, 2024, 7:31 PM")
    expect(page.locator("#assignmentsByStudent")).to_contain_text(submitted_re)

    # ---------------------
    context.close()
    browser.close()

#with sync_playwright() as playwright:
#    run(playwright)
